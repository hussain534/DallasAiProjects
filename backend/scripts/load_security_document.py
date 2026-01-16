"""
Script to load a DOCX document into security_items collection.

Usage:
    python scripts/load_security_document.py <docx_file_path> [document_number] [document_name]
"""

import asyncio
import sys
import json
from pathlib import Path
from docx import Document
from motor.motor_asyncio import AsyncIOMotorClient
from pymongo.errors import ConnectionFailure, ServerSelectionTimeoutError, DuplicateKeyError

# Add parent directory to path to import app modules
backend_dir = Path(__file__).parent.parent
sys.path.insert(0, str(backend_dir))

from app.core.config import settings
from app.utils.datetime_utils import utc_now


def extract_docx_content(file_path: str) -> dict:
    """
    Extract content from a DOCX file.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Dictionary with extracted content
    """
    print(f"Reading DOCX file: {file_path}")
    
    try:
        doc = Document(file_path)
        
        # Extract paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append({
                    "text": text,
                    "style": para.style.name if para.style else None
                })
        
        # Extract tables
        tables = []
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append(row_data)
            tables.append({
                "table_number": table_idx + 1,
                "rows": table_data
            })
        
        # Combine all text content
        full_text = "\n\n".join([p["text"] for p in paragraphs])
        
        content = {
            "file_name": Path(file_path).name,
            "paragraphs": paragraphs,
            "paragraph_count": len(paragraphs),
            "tables": tables,
            "table_count": len(tables),
            "full_text": full_text,
            "total_characters": len(full_text)
        }
        
        print(f"Extracted {len(paragraphs)} paragraphs and {len(tables)} tables")
        print(f"Total characters: {len(full_text)}")
        
        return content
        
    except Exception as e:
        print(f"Error reading DOCX file: {e}")
        import traceback
        traceback.print_exc()
        raise


async def upload_document_to_mongodb(
    docx_content: dict,
    document_number: int,
    document_name: str,
    collection_name: str = "security_items"
):
    """
    Upload document content to MongoDB security_items collection.
    
    Args:
        docx_content: Dictionary with extracted DOCX content
        document_number: Document number (must be unique)
        document_name: Document name
        collection_name: MongoDB collection name
    """
    import os
    connection_string = os.getenv("DATABASE_URL", settings.DATABASE_URL)
    database_name = os.getenv("DATABASE_NAME", settings.DATABASE_NAME)
    
    print(f"\nConnecting to MongoDB...")
    print(f"Database: {database_name}")
    print(f"Collection: {collection_name}")
    
    try:
        # Create MongoDB client
        client = AsyncIOMotorClient(
            connection_string,
            serverSelectionTimeoutMS=30000,
            connectTimeoutMS=30000,
            socketTimeoutMS=30000,
        )
        
        # Test connection
        await client.admin.command('ping')
        print("[OK] MongoDB connection successful")
        
        # Get database
        db = client[database_name]
        
        # Get collection
        collection = db[collection_name]
        
        # Check if document_number already exists
        existing = await collection.find_one({"document_number": document_number})
        if existing:
            print(f"\n[WARNING] Document with document_number={document_number} already exists")
            print(f"  Existing document_name: {existing.get('document_name', 'N/A')}")
            response = input("  Do you want to update it? (yes/no): ").strip().lower()
            if response != 'yes':
                print("[INFO] Operation cancelled")
                client.close()
                return False
            
            # Update existing document
            result = await collection.update_one(
                {"document_number": document_number},
                {
                    "$set": {
                        "document_name": document_name,
                        "document": docx_content,
                        "updated_at": utc_now()
                    }
                }
            )
            
            if result.modified_count > 0:
                print(f"[SUCCESS] Document updated successfully")
                print(f"  document_number: {document_number}")
                print(f"  document_name: {document_name}")
                print(f"  document size: {len(json.dumps(docx_content))} bytes")
            else:
                print(f"[WARNING] Document was not modified (may be identical)")
        else:
            # Insert new document
            document = {
                "document_number": document_number,
                "document_name": document_name,
                "document": docx_content,
                "created_at": utc_now(),
                "updated_at": utc_now()
            }
            
            try:
                await collection.insert_one(document)
                print(f"[SUCCESS] Document inserted successfully")
                print(f"  document_number: {document_number}")
                print(f"  document_name: {document_name}")
                print(f"  document size: {len(json.dumps(docx_content))} bytes")
            except DuplicateKeyError:
                print(f"[ERROR] Document with document_number={document_number} already exists (duplicate key)")
                client.close()
                return False
        
        # Verify insertion
        count = await collection.count_documents({})
        print(f"\n[INFO] Total documents in collection: {count}")
        
        client.close()
        return True
        
    except (ConnectionFailure, ServerSelectionTimeoutError) as e:
        print(f"[ERROR] Connection failed: {e}")
        return False
    except Exception as e:
        print(f"[ERROR] Error: {e}")
        import traceback
        traceback.print_exc()
        return False


async def main():
    """Main function."""
    # Parse command line arguments
    if len(sys.argv) < 2:
        print("Usage: python scripts/load_security_document.py <docx_file_path> [document_number] [document_name]")
        print("\nExample:")
        print('  python scripts/load_security_document.py "/path/to/file.docx" 1 "Security Framework"')
        sys.exit(1)
    
    file_path = sys.argv[1]
    document_number = int(sys.argv[2]) if len(sys.argv) > 2 else 1
    document_name = sys.argv[3] if len(sys.argv) > 3 else "Security Framework"
    
    # Check if file exists
    if not Path(file_path).exists():
        print(f"[ERROR] File not found: {file_path}")
        sys.exit(1)
    
    print("=" * 60)
    print("Load Security Document")
    print("=" * 60)
    print(f"File: {file_path}")
    print(f"Document Number: {document_number}")
    print(f"Document Name: {document_name}")
    print()
    
    # Extract content from DOCX
    try:
        docx_content = extract_docx_content(file_path)
    except Exception as e:
        print(f"[ERROR] Failed to extract content: {e}")
        sys.exit(1)
    
    # Upload to MongoDB
    success = await upload_document_to_mongodb(
        docx_content,
        document_number,
        document_name
    )
    
    if success:
        print("\n[SUCCESS] Operation completed")
        sys.exit(0)
    else:
        print("\n[ERROR] Operation failed")
        sys.exit(1)


if __name__ == "__main__":
    asyncio.run(main())

